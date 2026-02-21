from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

title = doc.add_heading('C语言程序代码摘要', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('项目概述', level=1)

files_info = [
    ("test.c", "迷宫路径搜索 (DFS)", "使用深度优先搜索算法在迷宫中寻找从起点S到终点P的路径。算法通过递归遍历四个方向，判断是否存在有效路径。"),
    ("1000.c", "两数求和", "基础输入输出程序，读取两个整数并输出它们的和。"),
    ("1001.c", "考试选择优化 (贪心算法)", "给定多个考试，每个考试有耗时和分值，在有限时间内选择考试以获得最高总分。采用价值/时间比作为贪心策略进行排序选择。"),
    ("1003.c", "鸡兔同笼问题", "经典的数学应用题，通过解方程组计算鸡和兔的数量。使用穷举法验证输入的有效性。"),
    ("1004.c", "赛车比赛策略 (贪心算法)", "判断赛车是否能赢得比赛。采用双指针贪心策略：最快对最快、最慢对最慢，尽可能赢得更多轮次。"),
    ("1005.c", "迷宫最短路径 (DFS/BFS)", "在迷宫中寻找从S到P的最短路径，要求步数不超过t。使用DFS遍历并记录所有路径的步数，取最小值判断。"),
    ("1005bfs.c", "BFS模板", "空的BFS算法模板文件。"),
    ("1007.c", "斐波那契数列", "递归计算斐波那契数列第n项。包含基础情况n=1返回1，n=2返回2。"),
    ("1008.c", "埃特巴什密码", "古典密码算法实现，大写字母和小写字母分别进行反向映射转换。"),
    ("1009.c", "最频繁元素频率", "找出数组中出现次数最多的元素的频次。使用排序后双指针统计各元素连续出现次数。"),
]

summary_text = f"""
本项目包含 {len([f for f in files_info if not f[0].startswith('1005bfs')])} 个C语言算法程序文件，主要涵盖以下算法类型：

1. 搜索算法：DFS/BFS迷宫路径搜索
2. 贪心算法：考试选择优化、赛车比赛策略
3. 递归算法：斐波那契数列
4. 密码学：埃特巴什古典密码
5. 统计分析：最频繁元素查找

这些程序均为算法竞赛或编程练习题目，每个程序独立解决特定问题。
"""

doc.add_paragraph(summary_text)

doc.add_heading('文件详细说明', level=1)

for filename, problem_name, description in files_info:
    doc.add_heading(f'{filename} - {problem_name}', level=2)
    doc.add_paragraph(description)

doc.add_heading('算法分类统计', level=1)

algo_types = {
    "搜索算法": ["test.c", "1005.c"],
    "贪心算法": ["1001.c", "1004.c"],
    "递归算法": ["1007.c"],
    "密码学": ["1008.c"],
    "基础算法": ["1000.c", "1003.c", "1009.c"],
    "未完成/模板": ["1005bfs.c"]
}

for algo_type, files in algo_types.items():
    doc.add_heading(algo_type, level=3)
    doc.add_paragraph('、'.join(files))

doc.add_heading('总结', level=1)

conclusion = """
本代码集展示了多种经典算法的C语言实现，包括：
- 深度优先搜索(DFS)和广度优先搜索(BFS)在迷宫问题中的应用
- 贪心策略在资源优化配置问题中的应用
- 递归思想在数列计算中的应用
- 古典密码算法的实现思路
- 排序和统计在实际问题中的应用

这些程序结构清晰，注释详细，适合作为算法学习和编程练习的参考资料。
"""

doc.add_paragraph(conclusion)

doc.save('代码摘要.docx')
print('Word文档已生成：代码摘要.docx')
